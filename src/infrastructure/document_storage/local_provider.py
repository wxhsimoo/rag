import os
import json
import hashlib
import base64
from typing import List, Dict, Any, Optional
import shutil
from pathlib import Path
from datetime import datetime

from .base import DocumentStorageProvider, RawDocument
from ...domain.entities.document import Document
from ...infrastructure.log.logger_service import LoggerService
from typing import List, Dict, Any, Optional
from datetime import datetime

class LocalDocumentStorageProvider(DocumentStorageProvider):
    """本地文件存储实现"""
    
    def __init__(self, data_path: str = None, logger: Optional[LoggerService] = None):
        self.data_path = Path(data_path)
        self.data_path.mkdir(parents=True, exist_ok=True)
        self.logger = logger
        # 维护文件ID与实际保存文件路径的映射文件路径
        self.id_map_path = self.data_path / 'id_map.json'

    def get_data_path(self) -> Path:
        """获取当前本地存储的数据根路径"""
        return self.data_path

    # --- 映射文件工具方法 ---
    def _load_id_map(self) -> Dict[str, Any]:
        try:
            if self.id_map_path.exists():
                with open(self.id_map_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return data if isinstance(data, dict) else {}
        except Exception as e:
            if self.logger:
                self.logger.warning(f"LocalStorage: load id_map failed: {e}")
        return {}

    def _save_id_map(self, id_map: Dict[str, Any]) -> None:
        try:
            with open(self.id_map_path, 'w', encoding='utf-8') as f:
                json.dump(id_map, f, ensure_ascii=False, indent=2)
        except Exception as e:
            if self.logger:
                self.logger.warning(f"LocalStorage: save id_map failed: {e}")

    def _get_mapped_path(self, document_id: str) -> Optional[Path]:
        id_map = self._load_id_map()
        entry = id_map.get(document_id)
        # 兼容旧结构：字符串为路径
        if isinstance(entry, str):
            return Path(entry) if entry else None
        # 新结构：字典包含 path 和 filename
        if isinstance(entry, dict):
            path = entry.get('path')
            return Path(path) if isinstance(path, str) and path else None
        return None

    def _set_mapped_path(self, document_id: str, path: Path) -> None:
        id_map = self._load_id_map()
        id_map[document_id] = {
            'path': str(path),
            'filename': path.name
        }
        self._save_id_map(id_map)

    def _remove_mapped_path(self, document_id: str) -> None:
        id_map = self._load_id_map()
        if document_id in id_map:
            del id_map[document_id]
            self._save_id_map(id_map)
    
    def _generate_document_id(self, content: str, source: str) -> str:
        """根据内容和来源生成文档ID"""
        content_hash = hashlib.md5(f"{source}:{content}".encode()).hexdigest()
        return f"doc_{content_hash[:16]}"
    
    def _load_json_file(self, file_path: Path) -> List[RawDocument]:
        """加载JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            documents = []
            if isinstance(data, list):
                for i, item in enumerate(data):
                    if isinstance(item, dict):
                        # 处理字典格式的数据
                        content = json.dumps(item, ensure_ascii=False, indent=2)
                        doc_id = self._generate_document_id(content, str(file_path))
                        documents.append(RawDocument(
                            id=f"{doc_id}_{i}",
                            content=content,
                            source=str(file_path),
                            metadata={"type": "json_object", "index": i}
                        ))
                    else:
                        # 处理其他类型的数据
                        content = str(item)
                        doc_id = self._generate_document_id(content, str(file_path))
                        documents.append(RawDocument(
                            id=f"{doc_id}_{i}",
                            content=content,
                            source=str(file_path),
                            metadata={"type": "json_item", "index": i}
                        ))
            elif isinstance(data, dict):
                # 处理单个字典
                content = json.dumps(data, ensure_ascii=False, indent=2)
                doc_id = self._generate_document_id(content, str(file_path))
                documents.append(RawDocument(
                    id=doc_id,
                    content=content,
                    source=str(file_path),
                    metadata={"type": "json_dict"}
                ))
            
            return documents
            
        except Exception as e:
            print(f"加载JSON文件失败 {file_path}: {str(e)}")
            return []
    
    def _load_text_file(self, file_path: Path) -> List[RawDocument]:
        """加载文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if not content.strip():
                return []
            
            doc_id = self._generate_document_id(content, str(file_path))
            return [RawDocument(
                id=doc_id,
                content=content,
                source=str(file_path),
                metadata={"type": "text_file", "size": len(content)}
            )]
            
        except Exception as e:
            print(f"加载文本文件失败 {file_path}: {str(e)}")
            return []
    
    def _split_markdown_sections(self, content: str, file_path: Path) -> List[RawDocument]:
        """将Markdown文件按章节分割"""
        documents = []
        sections = []
        current_section = []
        current_title = ""
        
        lines = content.split('\n')
        for line in lines:
            if line.startswith('#'):
                # 保存前一个章节
                if current_section:
                    section_content = '\n'.join(current_section).strip()
                    if section_content:
                        sections.append((current_title, section_content))
                
                # 开始新章节
                current_title = line.strip()
                current_section = [line]
            else:
                current_section.append(line)
        
        # 保存最后一个章节
        if current_section:
            section_content = '\n'.join(current_section).strip()
            if section_content:
                sections.append((current_title, section_content))
        
        # 创建文档
        for i, (title, section_content) in enumerate(sections):
            doc_id = self._generate_document_id(section_content, str(file_path))
            documents.append(RawDocument(
                id=f"{doc_id}_{i}",
                content=section_content,
                source=str(file_path),
                metadata={
                    "type": "markdown_section",
                    "title": title,
                    "section_index": i
                }
            ))
        
        return documents
    
    def _load_pdf_file(self, file_path: Path) -> List[RawDocument]:
        """加载PDF文件为文本"""
        try:
            try:
                from pdfminer.high_level import extract_text
                content = extract_text(str(file_path))
            except Exception:
                # 尝试使用 PyPDF2 作为备选
                import PyPDF2
                content = ""
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        content += page.extract_text() or ""
            if not content.strip():
                return []
            doc_id = self._generate_document_id(content, str(file_path))
            return [RawDocument(
                id=doc_id,
                content=content,
                source=str(file_path),
                metadata={"type": "pdf_file", "size": len(content)}
            )]
        except Exception as e:
            print(f"加载PDF文件失败 {file_path}: {str(e)}")
            return []

    def _load_docx_file(self, file_path: Path) -> List[RawDocument]:
        """加载Word(docx)文件为文本"""
        try:
            import docx
            d = docx.Document(str(file_path))
            paragraphs = [p.text for p in d.paragraphs if p.text]
            content = "\n".join(paragraphs)
            if not content.strip():
                return []
            doc_id = self._generate_document_id(content, str(file_path))
            return [RawDocument(
                id=doc_id,
                content=content,
                source=str(file_path),
                metadata={"type": "docx_file", "size": len(content)}
            )]
        except Exception as e:
            print(f"加载DOCX文件失败 {file_path}: {str(e)}")
            return []
    
    async def _load_raw_documents(self, source_path: Optional[str] = None) -> List[RawDocument]:
        """加载所有原始文档（可选按目录过滤）"""
        if self.logger:
            self.logger.info(f"LocalStorage: loading documents from {source_path or str(self.data_path)}")
        documents: List[RawDocument] = []
        base = self.data_path if not source_path else Path(source_path)
        for file_path in base.rglob('*'):
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                if suffix in ['.json']:
                    documents.extend(self._load_json_file(file_path))
                elif suffix in ['.md', '.markdown']:
                    # Markdown文件按章节分割
                    text_docs = self._load_text_file(file_path)
                    for doc in text_docs:
                        sections = self._split_markdown_sections(doc.content, file_path)
                        documents.extend(sections)
                elif suffix in ['.txt', '.text']:
                    documents.extend(self._load_text_file(file_path))
                elif suffix in ['.pdf']:
                    documents.extend(self._load_pdf_file(file_path))
                elif suffix in ['.docx']:
                    documents.extend(self._load_docx_file(file_path))
        if self.logger:
            self.logger.info(f"LocalStorage: loaded {len(documents)} documents")
        return documents
    
    async def _save_raw_document(self, document: RawDocument) -> bool:
        """保存原始文档到本地文件。
        规则：
        - 若 document.source 提供了文件名，则直接在目标目录以该文件名写入 document.content；
        - 否则，按默认逻辑将文本内容写入为 {document.id}.txt。
        不再生成任何 .meta.json 文件。
        """
        if self.logger:
            self.logger.info(f"LocalStorage: saving document {document.id}")
        try:
            # 若提供了源路径名，则直接写入内容到目标目录同名文件
            if document.source:
                src_path = Path(document.source)
                dst_path = self.data_path / src_path.name if src_path.name else (self.data_path / f"{document.id}.txt")
                # 若元数据包含二进制内容（base64），则按二进制写入，避免失真
                binary_b64 = None
                try:
                    binary_b64 = (document.metadata or {}).get('_binary_base64')
                except Exception:
                    binary_b64 = None
                if binary_b64:
                    try:
                        data = base64.b64decode(binary_b64)
                    except Exception:
                        data = b""
                    with open(dst_path, 'wb') as f:
                        f.write(data)
                else:
                    with open(dst_path, 'w', encoding='utf-8') as f:
                        f.write(document.content)
                # 更新映射
                self._set_mapped_path(document.id, dst_path)
                return True

            # 默认行为：写入为 {id}.txt
            file_path = self.data_path / f"{document.id}.txt"
            binary_b64 = None
            try:
                binary_b64 = (document.metadata or {}).get('_binary_base64')
            except Exception:
                binary_b64 = None
            if binary_b64:
                try:
                    data = base64.b64decode(binary_b64)
                except Exception:
                    data = b""
                with open(file_path, 'wb') as f:
                    f.write(data)
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(document.content)
            # 更新 id_map.json：记录 doc.id -> {id}.txt 路径
            # 更新映射
            self._set_mapped_path(document.id, file_path)
            return True
        except Exception as e:
            if self.logger:
                self.logger.error(f"LocalStorage: save failed {document.id}: {e}")
            else:
                print(f"保存文档失败: {str(e)}")
            return False
    
    async def _get_raw_document(self, document_id: str) -> Optional[RawDocument]:
        """读取原始文档：优先从 id_map.json 解析路径，其次回退到 {id}.txt。"""
        if self.logger:
            self.logger.info(f"LocalStorage: get document {document_id}")
        try:
            # 1) 从映射获取路径
            mapped_path: Optional[Path] = self._get_mapped_path(document_id)
            # 2) 若映射缺失，回退到 {id}.txt
            file_path = mapped_path if mapped_path is not None else (self.data_path / f"{document_id}.txt")
            if not file_path.exists():
                if self.logger:
                    self.logger.warning(f"LocalStorage: document not found {document_id}")
                return None
            # 3) 按后缀决定读取方式（文本/二进制）
            suffix = file_path.suffix.lower()
            if suffix in ['.txt', '.text', '.md', '.markdown']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            else:
                # 对于二进制类型，读取为字节并宽容解码
                with open(file_path, 'rb') as f:
                    content_bytes = f.read()
                content = content_bytes.decode('utf-8', errors='ignore')
            return RawDocument(id=document_id, content=content, source=str(file_path), metadata={})
        except Exception as e:
            if self.logger:
                self.logger.error(f"LocalStorage: get failed {document_id}: {e}")
            else:
                print(f"获取文档失败: {str(e)}")
            return None
    
    async def _delete_raw_document(self, document_id: str) -> bool:
        """删除原始文档"""
        if self.logger:
            self.logger.info(f"LocalStorage: delete document {document_id}")
        try:
            # 优先通过映射查找真实路径
            mapped_path = self._get_mapped_path(document_id)
            target_path = mapped_path if mapped_path is not None else (self.data_path / f"{document_id}.txt")
            success = True
            if target_path.exists():
                try:
                    target_path.unlink()
                except Exception:
                    success = False
            else:
                success = False
            # 删除成功后移除映射项
            if success:
                self._remove_mapped_path(document_id)
            return success
        except Exception as e:
            if self.logger:
                self.logger.error(f"LocalStorage: delete failed {document_id}: {e}")
            else:
                print(f"删除文档失败: {str(e)}")
            return False
    
    async def _list_raw_documents(self, category: Optional[str] = None) -> List[Dict[str, Any]]:
        """从 id_map.json 读取文档信息字典（id/source/metadata）。
        当映射文件不存在或不可读时，返回空列表并记录日志。
        """
        if self.logger:
            self.logger.info(f"LocalStorage: listing documents from id_map.json category={category}")
        infos: List[Dict[str, Any]] = []
        try:
            if not self.id_map_path.exists():
                if self.logger:
                    self.logger.warning(f"LocalStorage: id_map.json not found at {self.id_map_path}")
                return []
            with open(self.id_map_path, 'r', encoding='utf-8') as f:
                id_map = json.load(f)
            if not isinstance(id_map, dict):
                if self.logger:
                    self.logger.warning("LocalStorage: id_map.json content is not a dict")
                return []
            for doc_id, entry in id_map.items():
                # 兼容旧结构
                if isinstance(entry, str):
                    source = entry
                    filename = Path(entry).name
                elif isinstance(entry, dict):
                    source = entry.get('path') or ''
                    filename = entry.get('filename') or (Path(source).name if source else '')
                else:
                    continue
                if not isinstance(source, str):
                    continue
                if category and (category not in source):
                    continue
                file_path = Path(source)
                try:
                    size = file_path.stat().st_size if file_path.exists() else 0
                except Exception:
                    size = 0
                infos.append({
                    "id": doc_id,
                    "filename": filename,
                    "metadata": {"size": size}
                })
        except Exception as e:
            if self.logger:
                self.logger.error(f"LocalStorage: failed to list from id_map.json: {e}")
            else:
                print(f"列出文档失败: {str(e)}")
            return []
        if self.logger:
            self.logger.info(f"LocalStorage: listed {len(infos)} items from id_map.json")
        return infos
    
    def get_supported_formats(self) -> List[str]:
        """同步返回支持格式，保持与接口一致"""
        return ['.txt', '.text', '.md', '.markdown', '.pdf', '.docx']
    
    def _raw_to_domain_document(self, raw_doc: RawDocument) -> Document:
        """将RawDocument转换为Domain层Document"""
        return Document(
            content=raw_doc.content,
            metadata=raw_doc.metadata or {},
            doc_id=raw_doc.id,
            source_path=raw_doc.source,
            created_at=datetime.now()
        )
    
    def _domain_to_raw_document(self, doc: Document) -> RawDocument:
        """将Domain层Document转换为RawDocument"""
        try:
            print(f"LocalProvider: _domain_to_raw_document doc_id={doc.doc_id}")
        except Exception:
            pass
        return RawDocument(
            id=doc.doc_id or self._generate_document_id(doc.content, doc.source_path or ""),
            content=doc.content,
            source=doc.source_path or "",
            metadata=doc.metadata
        )
    
    # =====================
    # Domain层接口实现（符合 DocumentStorageService 签名）
    # =====================
    async def load_documents(self, source_path: Optional[str] = None) -> List[Document]:
        raw_docs = await self._load_raw_documents(source_path=source_path)
        return [self._raw_to_domain_document(d) for d in raw_docs]

    async def save_document(self, document: Document) -> tuple[bool, str]:
        raw_doc = self._domain_to_raw_document(document)
        ok = await self._save_raw_document(raw_doc)
        if ok and not document.doc_id:
            document.doc_id = raw_doc.id
        return ok, (raw_doc.id if ok else "")

    async def get_document(self, document_id: str) -> Optional[Document]:
        raw = await self._get_raw_document(document_id)
        return self._raw_to_domain_document(raw) if raw else None

    async def list_documents(self, category: Optional[str] = None) -> List[Dict[str, Any]]:
        return await self._list_raw_documents(category=category)

    async def delete_document(self, document_id: str) -> bool:
        return await self._delete_raw_document(document_id)

    async def search_documents(self, query: str, category: Optional[str] = None, limit: int = 10) -> List[Document]:
        all_docs = await self.load_documents()
        results: List[Document] = []
        for doc in all_docs:
            if query.lower() in (doc.content or "").lower():
                if category:
                    if (doc.metadata.get("category") != category) and (not (doc.source_path or "").startswith(category)):
                        continue
                results.append(doc)
                if len(results) >= limit:
                    break
        return results

    async def get_document_count(self, category: Optional[str] = None) -> int:
        infos = await self._list_raw_documents(category=category)
        return len(infos)

    async def update_document(self, document: Document) -> bool:
        raw_doc = self._domain_to_raw_document(document)
        return await self._save_raw_document(raw_doc)

    async def validate_document(self, document: Document) -> bool:
        if not (document.content or "").strip():
            return False
        if not document.doc_id:
            return False
        return True
