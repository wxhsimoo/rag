import asyncio
import hashlib
from typing import List, Dict, Any, Optional
from urllib.parse import urljoin

import boto3
from botocore.client import Config

from .base import DocumentStorageProvider, RawDocument
from ...infrastructure.log.logger_service import LoggerService
from ...domain.entities.document import Document
from datetime import datetime

class S3DocumentStorageProvider(DocumentStorageProvider):
    """S3 文档存储实现（兼容 MinIO / S3 协议）"""

    def __init__(
        self,
        bucket_name: str,
        endpoint_url: Optional[str] = None,
        access_key: Optional[str] = None,
        secret_key: Optional[str] = None,
        region_name: Optional[str] = None,
        use_ssl: bool = True,
        base_prefix: Optional[str] = None,
        logger: Optional[LoggerService] = None
    ):
        self.bucket = bucket_name
        self.prefix = base_prefix.strip("/") if base_prefix else ""
        self.logger = logger
        self.client = boto3.client(
            "s3",
            endpoint_url=endpoint_url,
            aws_access_key_id=access_key,
            aws_secret_access_key=secret_key,
            region_name=region_name,
            config=Config(s3={"addressing_style": "virtual"}, signature_version="s3v4"),
            use_ssl=use_ssl,
        )

    def _full_key(self, key: str) -> str:
        if not self.prefix:
            return key
        return f"{self.prefix}/{key}".lstrip("/")

    async def _list_objects(self, prefix: Optional[str] = None) -> List[Dict[str, Any]]:
        if self.logger:
            self.logger.info(f"S3Storage: list objects prefix={prefix or self.prefix}")
        def _sync_list():
            pfx = self._full_key(prefix) if prefix else self.prefix or ""
            paginator = self.client.get_paginator("list_objects_v2")
            pages = paginator.paginate(Bucket=self.bucket, Prefix=pfx)
            items: List[Dict[str, Any]] = []
            for page in pages:
                for obj in page.get("Contents", []):
                    items.append(obj)
            return items
        return await asyncio.to_thread(_sync_list)

    async def _load_raw_documents(self, source_path: Optional[str] = None) -> List[RawDocument]:
        """拉取对象并读取内容（注意：大规模对象可能较慢）"""
        if self.logger:
            self.logger.info(f"S3Storage: loading documents from prefix={source_path or self.prefix}")
        objects = await self._list_objects(prefix=source_path)
        docs: List[RawDocument] = []

        async def _get_one(key: str) -> Optional[RawDocument]:
            def _sync_get():
                resp = self.client.get_object(Bucket=self.bucket, Key=key)
                body_bytes = resp["Body"].read()
                meta = resp.get("Metadata", {}) or {}
                # 按后缀解析
                lower_key = key.lower()
                content = ""
                if lower_key.endswith(".pdf"):
                    try:
                        from pdfminer.high_level import extract_text
                        # pdfminer 不直接支持 bytes 流路径，先写到临时内存方案不直观；退回 PyPDF2
                        import PyPDF2
                        from io import BytesIO
                        reader = PyPDF2.PdfReader(BytesIO(body_bytes))
                        for page in reader.pages:
                            content += page.extract_text() or ""
                    except Exception:
                        content = body_bytes.decode("utf-8", errors="ignore")
                    doc_type = "pdf_file"
                elif lower_key.endswith(".docx"):
                    try:
                        import docx
                        from io import BytesIO
                        d = docx.Document(BytesIO(body_bytes))
                        paragraphs = [p.text for p in d.paragraphs if p.text]
                        content = "\n".join(paragraphs)
                    except Exception:
                        content = body_bytes.decode("utf-8", errors="ignore")
                    doc_type = "docx_file"
                else:
                    content = body_bytes.decode("utf-8", errors="ignore")
                    doc_type = "text_file"

                return RawDocument(
                    id=key.split("/")[-1],
                    content=content,
                    source=f"s3://{self.bucket}/{key}",
                    metadata={"type": doc_type, **meta},
                )
            return await asyncio.to_thread(_sync_get)

        for obj in objects:
            key = obj["Key"]
            try:
                rd = await _get_one(key)
                if rd:
                    docs.append(rd)
            except Exception as e:
                if self.logger:
                    self.logger.error(f"S3Storage: load failed key={key}: {e}")
                continue
        if self.logger:
            self.logger.info(f"S3Storage: loaded {len(docs)} documents")
        return docs

    async def _save_raw_document(self, document: RawDocument) -> bool:
        key = self._full_key(document.id)
        if self.logger:
            self.logger.info(f"S3Storage: saving document {document.id} to key={key}")
        def _sync_put():
            self.client.put_object(
                Bucket=self.bucket,
                Key=key,
                Body=(document.content or "").encode("utf-8"),
                Metadata=document.metadata or {},
                ContentType="text/plain; charset=utf-8",
            )
            return True
        try:
            return await asyncio.to_thread(_sync_put)
        except Exception as e:
            if self.logger:
                self.logger.error(f"S3Storage: save failed {document.id}: {e}")
            return False

    async def _get_raw_document(self, document_id: str) -> Optional[RawDocument]:
        key = self._full_key(document_id)
        if self.logger:
            self.logger.info(f"S3Storage: get document {document_id} from key={key}")
        def _sync_get():
            resp = self.client.get_object(Bucket=self.bucket, Key=key)
            body_bytes = resp["Body"].read()
            meta = resp.get("Metadata", {}) or {}
            lower_key = key.lower()
            content = ""
            if lower_key.endswith(".pdf"):
                try:
                    import PyPDF2
                    from io import BytesIO
                    reader = PyPDF2.PdfReader(BytesIO(body_bytes))
                    for page in reader.pages:
                        content += page.extract_text() or ""
                except Exception:
                    content = body_bytes.decode("utf-8", errors="ignore")
            elif lower_key.endswith(".docx"):
                try:
                    import docx
                    from io import BytesIO
                    d = docx.Document(BytesIO(body_bytes))
                    paragraphs = [p.text for p in d.paragraphs if p.text]
                    content = "\n".join(paragraphs)
                except Exception:
                    content = body_bytes.decode("utf-8", errors="ignore")
            else:
                content = body_bytes.decode("utf-8", errors="ignore")
            return RawDocument(
                id=document_id,
                content=content,
                source=f"s3://{self.bucket}/{key}",
                metadata=meta,
            )
        try:
            return await asyncio.to_thread(_sync_get)
        except Exception as e:
            if self.logger:
                self.logger.error(f"S3Storage: get failed {document_id}: {e}")
            return None

    async def _delete_raw_document(self, document_id: str) -> bool:
        key = self._full_key(document_id)
        if self.logger:
            self.logger.info(f"S3Storage: delete document {document_id} key={key}")
        def _sync_del():
            self.client.delete_object(Bucket=self.bucket, Key=key)
            return True
        try:
            return await asyncio.to_thread(_sync_del)
        except Exception as e:
            if self.logger:
                self.logger.error(f"S3Storage: delete failed {document_id}: {e}")
            return False

    async def _list_raw_documents(self, category: Optional[str] = None) -> List[Dict[str, Any]]:
        """列出对象（返回 id/source/metadata）"""
        prefix = category if category else None
        if self.logger:
            self.logger.info(f"S3Storage: listing documents prefix={prefix or self.prefix}")
        objects = await self._list_objects(prefix=prefix)
        infos: List[Dict[str, Any]] = []

        def _sync_head(key: str) -> Dict[str, Any]:
            try:
                head = self.client.head_object(Bucket=self.bucket, Key=key)
                meta = head.get("Metadata", {}) or {}
            except Exception:
                meta = {}
            return meta

        for obj in objects:
            key = obj["Key"]
            doc_id = key.split("/")[-1]
            meta = await asyncio.to_thread(_sync_head, key)
            infos.append({
                "id": doc_id,
                "source": f"s3://{self.bucket}/{key}",
                "metadata": meta,
            })
        if self.logger:
            self.logger.info(f"S3Storage: listed {len(infos)} items")
        return infos

    # =====================
    # Domain层接口实现（符合 DocumentStorageService 签名）
    # =====================
    def _generate_document_id(self, content: str, source: str) -> str:
        content_hash = hashlib.md5(f"{source}:{content}".encode()).hexdigest()
        return f"doc_{content_hash[:16]}"

    def _raw_to_domain_document(self, raw_doc: RawDocument) -> Document:
        return Document(
            content=raw_doc.content,
            metadata=raw_doc.metadata or {},
            doc_id=raw_doc.id,
            source_path=raw_doc.source,
            created_at=datetime.now()
        )

    def _domain_to_raw_document(self, doc: Document) -> RawDocument:
        return RawDocument(
            id=doc.doc_id or self._generate_document_id(doc.content, doc.source_path or ""),
            content=doc.content,
            source=doc.source_path or (doc.doc_id or ""),
            metadata=doc.metadata or {}
        )

    async def load_documents(self, source_path: Optional[str] = None) -> List[Document]:
        raw_docs = await self._load_raw_documents(source_path=source_path)
        return [self._raw_to_domain_document(d) for d in raw_docs]

    async def save_document(self, document: Document) -> tuple[bool, str]:
        raw_doc = self._domain_to_raw_document(document)
        ok = await self._save_raw_document(raw_doc)
        if ok and not document.doc_id:
            document.doc_id = raw_doc.id
        return ok, (raw_doc.id if ok else "")

    async def get_document(self, document_id: str) -> Optional[Document]:
        raw = await self._get_raw_document(document_id)
        return self._raw_to_domain_document(raw) if raw else None

    async def list_documents(self, category: Optional[str] = None) -> List[Dict[str, Any]]:
        return await self._list_raw_documents(category=category)

    async def delete_document(self, document_id: str) -> bool:
        return await self._delete_raw_document(document_id)

    async def search_documents(self, query: str, category: Optional[str] = None, limit: int = 10) -> List[Document]:
        all_docs = await self.load_documents()
        results: List[Document] = []
        for doc in all_docs:
            if query.lower() in (doc.content or "").lower():
                if category:
                    if (doc.metadata.get("category") != category) and (not (doc.source_path or "").startswith(category)):
                        continue
                results.append(doc)
                if len(results) >= limit:
                    break
        return results

    async def get_document_count(self, category: Optional[str] = None) -> int:
        infos = await self._list_raw_documents(category=category)
        return len(infos)

    async def update_document(self, document: Document) -> bool:
        raw_doc = self._domain_to_raw_document(document)
        return await self._save_raw_document(raw_doc)

    def get_supported_formats(self) -> List[str]:
        return ['.txt', '.text', '.md', '.markdown', '.pdf', '.docx']

    async def validate_document(self, document: Document) -> bool:
        return bool(document.doc_id and (document.content or "").strip())
