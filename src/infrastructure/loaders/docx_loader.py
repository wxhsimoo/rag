from typing import List, Dict, Any, Optional

try:
    from langchain_community.document_loaders import UnstructuredWordDocumentLoader
except ImportError:
    UnstructuredWordDocumentLoader = None

from ...infrastructure.log.logger_service import LoggerService
from ..splitters.types import InfraDocument
from .base import InfraDocumentLoader


class DocxDocumentLoader(InfraDocumentLoader):
    """DOCX 文档加载器

    优先使用 langchain 的 `UnstructuredWordDocumentLoader`，不可用时回退到 `python-docx`。
    """

    def __init__(self, logger: Optional[LoggerService] = None):
        self.logger = logger

    def supports_file_type(self, file_path: str) -> bool:
        return file_path.lower().endswith(".docx")

    def get_supported_extensions(self) -> List[str]:
        return ["docx"]

    def load(self, file_path: str) -> List[InfraDocument]:
        return self._load_documents(file_path)

    def _create_document(self, content: str, metadata: Dict[str, Any] = None) -> InfraDocument:
        if metadata is None:
            metadata = {}
        return InfraDocument(content=content, metadata=metadata)

    def _clean_content(self, content: str) -> str:
        if not content:
            return ""
        content = content.strip()
        content = content.replace("\r\n", "\n").replace("\r", "\n")
        return content

    def _load_documents(self, file_path: str) -> List[InfraDocument]:
        if UnstructuredWordDocumentLoader is not None:
            return self._load_with_langchain(file_path)
        return self._load_with_builtin(file_path)

    def _load_with_langchain(self, file_path: str) -> List[InfraDocument]:
        try:
            loader = UnstructuredWordDocumentLoader(file_path)
            lc_docs = loader.load()
            documents: List[Document] = []
            for lc in lc_docs:
                metadata = dict(lc.metadata) if lc.metadata else {}
                doc = self._create_document(content=self._clean_content(lc.page_content), metadata=metadata)
                documents.append(doc)
            return documents
        except Exception as e:
            if self.logger:
                self.logger.warning(f"使用 UnstructuredWordDocumentLoader 加载失败，回退到内置实现: {e}")
            return self._load_with_builtin(file_path)

    def _load_with_builtin(self, file_path: str) -> List[InfraDocument]:
        try:
            import docx  # python-docx
            d = docx.Document(str(file_path))
            paragraphs = [p.text for p in d.paragraphs if p.text]
            content = "\n".join(paragraphs)
            content = self._clean_content(content)
            if not content:
                return []
            meta = {
                "type": "docx_file",
                "source": str(file_path),
                "size": len(content),
            }
            return [self._create_document(content=content, metadata=meta)]
        except Exception as e:
            raise ValueError(f"加载DOCX文件失败: {e}")
